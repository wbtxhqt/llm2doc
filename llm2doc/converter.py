from __future__ import annotations

import datetime
import json
import os
from typing import Any, Callable, Dict, List, Optional, Union

from docx import Document
from docx.enum.dml import MSO_COLOR_TYPE, MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_COLOR_INDEX,
    WD_LINE_SPACING,
    WD_UNDERLINE,
)
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run as _Run
# Namespaces used when inspecting low-level XML
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}

_id_counter = 0


def _get_next_id():
    global _id_counter
    _id_counter += 1
    return f"doc-obj-{_id_counter}"


def _len_to_pt(x) -> Optional[float]:
    if x is None:
        return None
    try:
        return float(x.pt)
    except Exception:
        try:
            # Sometimes values are stored as Emu or Twips lengths in python-docx,
            # but Length exposes .pt. If not, try direct cast.
            return float(x)
        except Exception:
            return None

def _color_to_hex(color) -> Optional[str]:
    if color is None or not hasattr(color, "type") or color.type is None:
        return None
    try:
        if color.type == MSO_COLOR_TYPE.RGB:
            if color.rgb is not None:
                return f"#{color.rgb}"
        elif color.type == MSO_COLOR_TYPE.THEME:
            return color.theme_color.name
        elif color.type == MSO_COLOR_TYPE.AUTO:
            return "#000000"
    except Exception:
        pass
    return None

def _underline_value(u) -> Optional[Union[bool, str]]:
    # python-docx run.underline may be True/False/None or a WD_UNDERLINE enum.
    if u is None or isinstance(u, bool):
        return u
    try:
        # Enum value -> its name (e.g., 'SINGLE', 'DOUBLE', ...)
        return u.name if hasattr(u, "name") else str(u)
    except Exception:
        return str(u)

def _line_spacing_info(pf) -> Dict[str, Optional[Union[float, str]]]:
    # line_spacing can be a multiple (float) or an absolute length
    ls_mult = None
    ls_pt = None
    try:
        ls = pf.line_spacing
        if isinstance(ls, (int, float)):
            ls_mult = float(ls)
        else:
            ls_pt = _len_to_pt(ls)
    except Exception:
        pass
    try:
        rule = pf.line_spacing_rule.name if pf.line_spacing_rule else None
    except Exception:
        rule = None
    return {
        "line_spacing_multiple": ls_mult,
        "line_spacing_pt": ls_pt,
        "line_spacing_rule": rule,
    }

def _get_paragraph_numbering_info(doc: Document, p: Paragraph) -> Optional[Dict[str, Any]]:
    # Detect Word numbering (bullets/numbered lists)
    try:
        pPr = p._p.pPr
        if pPr is None or pPr.numPr is None:
            return None
        num_id = None
        ilvl = 0
        if pPr.numPr.numId is not None and pPr.numPr.numId.val is not None:
            num_id = int(pPr.numPr.numId.val)
        if pPr.numPr.ilvl is not None and pPr.numPr.ilvl.val is not None:
            ilvl = int(pPr.numPr.ilvl.val)
        result = {"numId": num_id, "level": ilvl, "format": None, "lvlText": None}

        # Try to resolve bullet/number format from numbering part
        numbering_part = getattr(doc.part, "numbering_part", None)
        if numbering_part is None or num_id is None:
            return result

        ne = numbering_part.element  # CT_Numbering element
        ns = dict(ne.nsmap)
        ns.update(NS)

        abs_nodes = ne.xpath(f'.//w:num[@w:numId="{num_id}"]/w:abstractNumId', namespaces=ns)
        if not abs_nodes:
            return result
        abs_id = abs_nodes[0].get(qn("w:val"))

        lvl_nodes = ne.xpath(
            f'.//w:abstractNum[@w:abstractNumId="{abs_id}"]/w:lvl[@w:ilvl="{ilvl}"]',
            namespaces=ns,
        )
        if not lvl_nodes:
            return result

        fmt_node = lvl_nodes[0].find(qn("w:numFmt"))
        txt_node = lvl_nodes[0].find(qn("w:lvlText"))
        if fmt_node is not None:
            result["format"] = fmt_node.get(qn("w:val"))
        if txt_node is not None:
            result["lvlText"] = txt_node.get(qn("w:val"))
        return result
    except Exception:
        return None

def _get_hyperlink_info(run) -> Optional[Dict[str, Any]]:
    try:
        r = run._r
        parent = r.getparent()
        # climb up until hyperlink or root
        while parent is not None and parent.tag != qn("w:hyperlink"):
            parent = parent.getparent()
        if parent is None:
            return None
        rId = parent.get(qn("r:id"))
        anchor = parent.get(qn("w:anchor"))
        url = None
        if rId and rId in run.part.rels:
            rel = run.part.rels[rId]
            try:
                # External link
                url = str(rel.target_ref)
            except Exception:
                try:
                    url = str(rel._target)  # fallback
                except Exception:
                    url = None
        return {"rId": rId, "url": url, "anchor": anchor}
    except Exception:
        return None

def _get_run_images(run) -> List[Dict[str, Any]]:
    # Detect images embedded in this run (inline shapes)
    images = []
    try:
        # Look for DrawingML blips
        blips = run._r.xpath(".//a:blip", namespaces=NS)
        for blip in blips:
            rId = blip.get(qn("r:embed"))
            if not rId:
                continue
            rel = run.part.rels.get(rId)
            if rel is None:
                continue
            # For images, target_part has .image with filename, content_type
            part = getattr(rel, "target_part", None)
            if part is None:
                # some versions expose _target_part
                part = getattr(rel, "_target_part", None)
            info = {"rId": rId, "filename": None, "content_type": None}
            try:
                img = getattr(part, "image", None)
                if img is not None:
                    info["filename"] = img.filename
                    info["content_type"] = getattr(part, "content_type", None)
                else:
                    # fallback info
                    info["filename"] = getattr(part, "partname", None)
                    info["content_type"] = getattr(part, "content_type", None)
            except Exception:
                pass
            images.append(info)
    except Exception:
        pass
    return images

def _serialize_run(run) -> Dict[str, Any]:
    font = run.font
    highlight = None
    try:
        highlight = font.highlight_color.name if font.highlight_color else None
    except Exception:
        highlight = None

    run_obj: Dict[str, Any] = {
        "id": _get_next_id(),
        "type": "run",
        "text": run.text,
        "style": run.style.name if getattr(run, "style", None) else None,
        "bold": bool(run.bold) if run.bold is not None else None,
        "italic": bool(run.italic) if run.italic is not None else None,
        "underline": _underline_value(run.underline),
        "font": {
            "name": font.name,
            "size_pt": _len_to_pt(font.size),
            "color": _color_to_hex(font.color),
            "highlight": highlight,
            "all_caps": bool(font.all_caps) if font.all_caps is not None else None,
            "small_caps": bool(font.small_caps) if font.small_caps is not None else None,
            "strike": bool(font.strike) if font.strike is not None else None,
            "double_strike": bool(font.double_strike) if font.double_strike is not None else None,
            "superscript": bool(font.superscript) if font.superscript is not None else None,
            "subscript": bool(font.subscript) if font.subscript is not None else None,
        },
        "hyperlink": _get_hyperlink_info(run),
        "images": _get_run_images(run) or None,
    }

    # If run contains only images and no text, keep text as empty string; that's fine.
    return run_obj

def _serialize_paragraph(doc: Document, p: Paragraph, compact: bool = False) -> Dict[str, Any]:
    pf = p.paragraph_format
    align = None
    try:
        align = p.alignment.name if p.alignment else None
    except Exception:
        align = None

    runs_data = [_serialize_run(r) for r in p.runs]
    if compact:
        runs_data = _merge_runs(runs_data)

    para_obj: Dict[str, Any] = {
        "id": _get_next_id(),
        "type": "paragraph",
        "style": p.style.name if getattr(p, "style", None) else None,
        "alignment": align,
        "numbering": _get_paragraph_numbering_info(doc, p),
        "paragraph_format": {
            "left_indent_pt": _len_to_pt(getattr(pf, "left_indent", None)),
            "right_indent_pt": _len_to_pt(getattr(pf, "right_indent", None)),
            "first_line_indent_pt": _len_to_pt(getattr(pf, "first_line_indent", None)),
            "space_before_pt": _len_to_pt(getattr(pf, "space_before", None)),
            "space_after_pt": _len_to_pt(getattr(pf, "space_after", None)),
            **_line_spacing_info(pf),
        },
        "runs": runs_data,
    }
    return para_obj

def _iter_block_items(parent) -> Any:
    # Preserve document order of paragraphs and tables
    from docx.document import Document as _Document
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("Unsupported parent for block iteration")

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)

def _serialize_cell(cell: _Cell, doc: Document, compact: bool = False) -> Dict[str, Any]:
    # Collect the blocks within a cell (paragraphs and nested tables)
    blocks: List[Dict[str, Any]] = []
    for item in _iter_block_items(cell):
        if isinstance(item, Paragraph):
            blocks.append(_serialize_paragraph(doc, item, compact=compact))
        elif isinstance(item, Table):
            blocks.append(_serialize_table(doc, item, compact=compact))
    v_align = None
    try:
        v_align = cell.vertical_alignment.name if cell.vertical_alignment else None
    except Exception:
        v_align = None
    return {
        "id": _get_next_id(),
        "type": "cell",
        "vertical_alignment": v_align,
        "blocks": blocks,
    }

def _serialize_table(doc: Document, tbl: Table, compact: bool = False) -> Dict[str, Any]:
    table_obj: Dict[str, Any] = {
        "id": _get_next_id(),
        "type": "table",
        "style": tbl.style.name if getattr(tbl, "style", None) else None,
        "rows": [],
    }
    for row in tbl.rows:
        row_cells = [_serialize_cell(cell, doc, compact=compact) for cell in row.cells]
        table_obj["rows"].append(row_cells)
    return table_obj

def _core_properties(doc: Document) -> Dict[str, Any]:
    cp = doc.core_properties
    def ts(x):
        if not x:
            return None
        if isinstance(x, datetime.datetime):
            return x.isoformat()
        return str(x)
    return {
        "title": cp.title,
        "subject": cp.subject,
        "category": cp.category,
        "keywords": cp.keywords,
        "comments": cp.comments,
        "author": cp.author,
        "last_modified_by": cp.last_modified_by,
        "created": ts(cp.created),
        "modified": ts(cp.modified),
        "version": cp.version,
    }

def _section_info(doc: Document) -> List[Dict[str, Any]]:
    infos = []
    for s in doc.sections:
        try:
            infos.append({
                "page_width_pt": _len_to_pt(s.page_width),
                "page_height_pt": _len_to_pt(s.page_height),
                "left_margin_pt": _len_to_pt(s.left_margin),
                "right_margin_pt": _len_to_pt(s.right_margin),
                "top_margin_pt": _len_to_pt(s.top_margin),
                "bottom_margin_pt": _len_to_pt(s.bottom_margin),
                "header_distance_pt": _len_to_pt(s.header_distance),
                "footer_distance_pt": _len_to_pt(s.footer_distance),
                "orientation": s.orientation.name if s.orientation else None,
            })
        except Exception:
            infos.append({})
    return infos

def docx2json(
    docx_path: str,
    json_path: Optional[str] = None,
    indent: int = 2,
    compact: bool = False,
) -> Dict[str, Any]:
    """
    Convert a .docx file into a JSON-serializable Python dict describing the document
    for downstream LLM editing of content and formatting.

    Args:
        docx_path: Path to the input .docx file.
        json_path: Optional path to write a JSON file. If None, no file is written.
        indent: JSON indentation when writing to file.
        compact: If True, remove null values and merge runs to reduce size.

    Returns:
        A dict with:
        - meta: basic info and core properties
        - sections: page layout info
        - blocks: ordered list of top-level blocks (paragraph or table)
    """
    global _id_counter
    _id_counter = 0
    doc = Document(docx_path)

    blocks: List[Dict[str, Any]] = []
    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            blocks.append(_serialize_paragraph(doc, item, compact=compact))
        elif isinstance(item, Table):
            blocks.append(_serialize_table(doc, item, compact=compact))

    result: Dict[str, Any] = {
        "version": "1.0",
        "meta": {
            "source": docx_path,
            "generated_at": datetime.datetime.utcnow().isoformat() + "Z",
            "core_properties": _core_properties(doc),
        },
        "sections": _section_info(doc),
        "blocks": blocks,
    }
    if compact:
        result = _clean_dict(result)

    if json_path:
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=indent)

    return result

# Example usage:
# data = docx2json("input.docx", "output.json")
# If you just want the dict without writing a file:
# data = docx2json("input.docx")



def _clean_dict(d: Any) -> Any:
    """Recursively remove None values, empty strings, and empty lists/dicts."""
    if isinstance(d, dict):
        return {k: _clean_dict(v) for k, v in d.items() if v is not None and v != "" and v != [] and v != {}}
    if isinstance(d, list):
        return [_clean_dict(i) for i in d if i is not None and i != "" and i != [] and i != {}]
    return d

def _merge_runs(runs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Merge consecutive runs with identical styling."""
    if not runs:
        return []
    merged = []
    current_run = runs[0]
    for next_run in runs[1:]:
        # Merge if text is present and styles are identical
        if current_run.get("text") and next_run.get("text") and _are_styles_equal(current_run, next_run):
            current_run["text"] += next_run["text"]
        else:
            merged.append(current_run)
            current_run = next_run
    merged.append(current_run)
    return merged

def _are_styles_equal(run1: Dict[str, Any], run2: Dict[str, Any]) -> bool:
    """Check if two run objects have the same styling properties after cleaning."""
    style1 = {k: v for k, v in run1.items() if k != "text"}
    style2 = {k: v for k, v in run2.items() if k != "text"}
    return _clean_dict(style1) == _clean_dict(style2)

def _pt(x: Optional[float]):
    return Pt(float(x)) if isinstance(x, (int, float)) else None

def _rgb_from_hex(s: Optional[str]) -> Optional[RGBColor]:
    if not s:
        return None
    s = s.strip()
    if s.startswith("#"):
        s = s[1:]
    if len(s) != 6:
        return None
    try:
        r = int(s[0:2], 16)
        g = int(s[2:4], 16)
        b = int(s[4:6], 16)
        return RGBColor(r, g, b)
    except Exception:
        return None

def _get_enum(enum_cls, name: Optional[str]):
    if not name:
        return None
    try:
        return getattr(enum_cls, str(name).upper())
    except Exception:
        # sometimes name is already an enum, or mixed case; try case-insensitive
        try:
            vals = {k.upper(): v for k, v in enum_cls.__members__.items()}  # type: ignore
            return vals.get(str(name).upper())
        except Exception:
            return None

def _safe_set_style(obj, style_name: Optional[str]):
    if not style_name:
        return
    try:
        obj.style = style_name
    except Exception:
        # Unknown style in this document template; ignore.
        pass

def _apply_paragraph_format(p: Paragraph, fmt: Dict[str, Any], alignment: Optional[str]):
    if alignment:
        enum = _get_enum(WD_ALIGN_PARAGRAPH, alignment)
        if enum:
            p.alignment = enum

    pf = p.paragraph_format
    if not isinstance(fmt, dict):
        return

    if (v := fmt.get("left_indent_pt")) is not None:
        pf.left_indent = _pt(v)
    if (v := fmt.get("right_indent_pt")) is not None:
        pf.right_indent = _pt(v)
    if (v := fmt.get("first_line_indent_pt")) is not None:
        pf.first_line_indent = _pt(v)
    if (v := fmt.get("space_before_pt")) is not None:
        pf.space_before = _pt(v)
    if (v := fmt.get("space_after_pt")) is not None:
        pf.space_after = _pt(v)

    # line-spacing: multiple or absolute length
    ls_mult = fmt.get("line_spacing_multiple")
    ls_pt = fmt.get("line_spacing_pt")
    if isinstance(ls_mult, (int, float)):
        pf.line_spacing = float(ls_mult)
    elif isinstance(ls_pt, (int, float)):
        pf.line_spacing = _pt(ls_pt)

    rule = fmt.get("line_spacing_rule")
    enum = _get_enum(WD_LINE_SPACING, rule)
    if enum:
        pf.line_spacing_rule = enum

def _underline_from_json(u: Any):
    # Accept True/False/None or strings like 'SINGLE', 'DOUBLE', 'NONE'
    if u is None or isinstance(u, bool):
        return u
    enum = _get_enum(WD_UNDERLINE, str(u))
    if enum:
        return enum
    # fallback to True for any non-empty truthy value
    return True

def _apply_run_formatting(run: _Run, run_obj: Dict[str, Any]):
    # style
    _safe_set_style(run, run_obj.get("style"))
    # basic toggles
    if run_obj.get("bold") is not None:
        run.bold = bool(run_obj.get("bold"))
    if run_obj.get("italic") is not None:
        run.italic = bool(run_obj.get("italic"))

    # underline (bool or enum name)
    if "underline" in run_obj:
        run.underline = _underline_from_json(run_obj.get("underline"))

    font = run.font
    fobj = run_obj.get("font") or {}

    if (name := fobj.get("name")):
        font.name = name
    if (size := fobj.get("size_pt")) is not None:
        font.size = _pt(size)
    if (col := fobj.get("color")):
        if col.startswith("#"):
            rgb = _rgb_from_hex(col)
            if rgb:
                font.color.rgb = rgb
        else:
            try:
                theme_color_idx = getattr(MSO_THEME_COLOR_INDEX, col)
                font.color.theme_color = theme_color_idx
            except (AttributeError, TypeError):
                pass

    if (hi := fobj.get("highlight")):
        enum = _get_enum(WD_COLOR_INDEX, hi)
        if enum:
            font.highlight_color = enum

    def set_bool(attr, key):
        val = fobj.get(key)
        if val is not None:
            setattr(font, attr, bool(val))
    set_bool("all_caps", "all_caps")
    set_bool("small_caps", "small_caps")
    set_bool("strike", "strike")
    set_bool("double_strike", "double_strike")
    set_bool("superscript", "superscript")
    set_bool("subscript", "subscript")

def _add_hyperlink_run(paragraph: Paragraph, text: str, url: Optional[str], anchor: Optional[str]) -> _Run:
    """
    Create a clickable hyperlink run in the paragraph.
    If url is provided, creates an external hyperlink.
    If anchor is provided (bookmark), creates an internal link.
    Returns a python-docx Run object for further formatting.
    """
    # Build <w:hyperlink>
    hyperlink = OxmlElement("w:hyperlink")
    part = paragraph.part

    if url:
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), r_id)
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)

    # Create the run element inside hyperlink
    new_run = OxmlElement("w:r")
    new_rPr = OxmlElement("w:rPr")
    new_run.append(new_rPr)
    t = OxmlElement("w:t")
    # preserve spaces
    t.set(qn("xml:space"), "preserve")
    t.text = text or ""
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    # Wrap into a Run for formatting
    return _Run(new_run, paragraph)

def _apply_numbering_style_if_any(p: Paragraph, numbering: Optional[Dict[str, Any]], paragraph_style_present: bool):
    """
    Approximate Word numbering using built-in list styles.
    If paragraph already has a style set, we don't override it.
    """
    if not numbering or paragraph_style_present:
        return
    fmt = (numbering.get("format") or "").lower()
    level = int(numbering.get("level") or 0)

    # Choose a base style name
    if fmt == "bullet":
        base = "List Bullet"
    else:
        # treat anything else as decimal-numbered
        base = "List Number"

    style_name = base if level <= 0 else f"{base} {level+1}"
    _safe_set_style(p, style_name)

def _default_image_resolver(resources_dir: Optional[str], info: Dict[str, Any]) -> Optional[str]:
    filename = info.get("filename")
    if not filename or not resources_dir:
        return None
    p = os.path.join(resources_dir, filename)
    return p if os.path.isfile(p) else None

def _add_images_for_run(run: _Run, images: Optional[List[Dict[str, Any]]],
                        image_resolver: Optional[Callable[[Dict[str, Any]], Optional[str]]],
                        resources_dir: Optional[str]):
    if not images:
        return
    for img in images:
        path = None
        if image_resolver:
            try:
                path = image_resolver(img)
            except Exception:
                path = None
        if not path:
            path = _default_image_resolver(resources_dir, img)

        if path and os.path.isfile(path):
            try:
                run.add_picture(path)
            except Exception:
                # fallback to a placeholder text if something goes wrong
                run.add_text(f"[Image]")
        else:
            # Can't resolve image: add a placeholder text to keep a clue in output
            if (label := img.get("filename")):
                run.add_text(f"[Image: {label}]")
            else:
                run.add_text("[Image]")

def _write_paragraph(doc: Document,
                     block: Dict[str, Any],
                     image_resolver: Optional[Callable[[Dict[str, Any]], Optional[str]]],
                     resources_dir: Optional[str]) -> Optional[Paragraph]:
    # Don't render a paragraph if it has no content, as it just adds a blank line.
    # A paragraph is considered empty if it has no runs or all its runs are empty (no text/images).
    is_empty = not any(r.get("text") or r.get("images") for r in block.get("runs", []))
    if is_empty:
        return None

    p = doc.add_paragraph()
    # paragraph style and format
    style_name = block.get("style")
    _safe_set_style(p, style_name)
    _apply_paragraph_format(p, block.get("paragraph_format") or {}, block.get("alignment"))

    # numbering approximation (if style not explicitly set)
    _apply_numbering_style_if_any(p, block.get("numbering"), paragraph_style_present=bool(style_name))

    # runs
    for r in block.get("runs", []):
        text = r.get("text") or ""
        hyperlink = r.get("hyperlink") or {}
        url = hyperlink.get("url")
        anchor = hyperlink.get("anchor")

        if url or anchor:
            run = _add_hyperlink_run(p, text, url, anchor)
        else:
            run = p.add_run(text)

        _apply_run_formatting(run, r)
        _add_images_for_run(run, r.get("images"), image_resolver, resources_dir)

    return p

def _write_table(doc: Document,
                 block: Dict[str, Any],
                 image_resolver: Optional[Callable[[Dict[str, Any]], Optional[str]]],
                 resources_dir: Optional[str]):
    rows = block.get("rows") or []
    n_rows = len(rows)
    n_cols = max((len(r) for r in rows), default=0)
    if n_rows == 0 or n_cols == 0:
        return

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    _safe_set_style(tbl, block.get("style"))

    for i, row in enumerate(rows):
        for j, cell_obj in enumerate(row):
            cell = tbl.cell(i, j)
            # vertical alignment
            v_align = _get_enum(WD_CELL_VERTICAL_ALIGNMENT, cell_obj.get("vertical_alignment"))
            if v_align:
                cell.vertical_alignment = v_align

            blocks = cell_obj.get("blocks") or []

            # Reuse the default empty paragraph for the first paragraph block if possible
            def write_block_into_cell(_block: Dict[str, Any]):
                btype = _block.get("type")
                if btype == "paragraph":
                    # Don't render a paragraph if it has no content.
                    is_empty = not any(r.get("text") or r.get("images") for r in _block.get("runs", []))
                    if is_empty:
                        return

                    # If the cell currently contains exactly one empty paragraph, reuse it
                    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text and not cell.paragraphs[0].runs:
                        p = cell.paragraphs[0]
                        # Apply style/format to reused paragraph
                        style_name = _block.get("style")
                        _safe_set_style(p, style_name)
                        _apply_paragraph_format(p, _block.get("paragraph_format") or {}, _block.get("alignment"))
                        _apply_numbering_style_if_any(p, _block.get("numbering"), paragraph_style_present=bool(style_name))

                        for r in _block.get("runs", []):
                            text = r.get("text") or ""
                            hyperlink = r.get("hyperlink") or {}
                            url = hyperlink.get("url")
                            anchor = hyperlink.get("anchor")
                            if url or anchor:
                                run = _add_hyperlink_run(p, text, url, anchor)
                            else:
                                run = p.add_run(text)
                            _apply_run_formatting(run, r)
                            _add_images_for_run(run, r.get("images"), image_resolver, resources_dir)
                    else:
                        # Add a fresh paragraph
                        p = cell.add_paragraph()
                        _safe_set_style(p, _block.get("style"))
                        _apply_paragraph_format(p, _block.get("paragraph_format") or {}, _block.get("alignment"))
                        _apply_numbering_style_if_any(p, _block.get("numbering"), paragraph_style_present=bool(_block.get("style")))
                        for r in _block.get("runs", []):
                            text = r.get("text") or ""
                            hyperlink = r.get("hyperlink") or {}
                            url = hyperlink.get("url")
                            anchor = hyperlink.get("anchor")
                            if url or anchor:
                                run = _add_hyperlink_run(p, text, url, anchor)
                            else:
                                run = p.add_run(text)
                            _apply_run_formatting(run, r)
                            _add_images_for_run(run, r.get("images"), image_resolver, resources_dir)
                elif btype == "table":
                    # nested table
                    sub_rows = _block.get("rows") or []
                    rcount = len(sub_rows)
                    ccount = max((len(sr) for sr in sub_rows), default=0)
                    if rcount and ccount:
                        sub_tbl = cell.add_table(rows=rcount, cols=ccount)
                        _safe_set_style(sub_tbl, _block.get("style"))
                        for ii, sr in enumerate(sub_rows):
                            for jj, sc in enumerate(sr):
                                sub_cell = sub_tbl.cell(ii, jj)
                                v_align2 = _get_enum(WD_CELL_VERTICAL_ALIGNMENT, sc.get("vertical_alignment"))
                                if v_align2:
                                    sub_cell.vertical_alignment = v_align2
                                for sb in sc.get("blocks") or []:
                                    write_block_into_cell(sb)

            for b in blocks:
                write_block_into_cell(b)

def _apply_core_properties(doc: Document, props: Dict[str, Any]):
    if not isinstance(props, dict):
        return
    cp = doc.core_properties
    def set_if_present(attr, key):
        val = props.get(key)
        if val is not None:
            try:
                setattr(cp, attr, val)
            except Exception:
                pass
    set_if_present("title", "title")
    set_if_present("subject", "subject")
    set_if_present("category", "category")
    set_if_present("keywords", "keywords")
    set_if_present("comments", "comments")
    set_if_present("author", "author")
    set_if_present("last_modified_by", "last_modified_by")
    # created/modified are handled by Word; setting strings may raise; skip.

def _apply_section_settings(doc: Document, sections_info: List[Dict[str, Any]]):
    if not sections_info:
        return
    # Apply only to the first section (we don't have block-to-section boundaries)
    s = doc.sections[0]
    info = sections_info[0] or {}
    def set_len(attr, key):
        v = info.get(key)
        if isinstance(v, (int, float)):
            try:
                setattr(s, attr, _pt(v))
            except Exception:
                pass
    set_len("page_width", "page_width_pt")
    set_len("page_height", "page_height_pt")
    set_len("left_margin", "left_margin_pt")
    set_len("right_margin", "right_margin_pt")
    set_len("top_margin", "top_margin_pt")
    set_len("bottom_margin", "bottom_margin_pt")
    set_len("header_distance", "header_distance_pt")
    set_len("footer_distance", "footer_distance_pt")
    # orientation enum might be present, but python-docx auto-sets with page sizes;
    # if present, we can try to set it.
    try:
        from docx.enum.section import WD_ORIENTATION
        orient = _get_enum(WD_ORIENTATION, info.get("orientation"))
        if orient:
            s.orientation = orient
    except Exception:
        pass

def json2docx(
    json_source: Union[str, Dict[str, Any]],
    output_docx_path: str,
    resources_dir: Optional[str] = None,
    image_resolver: Optional[Callable[[Dict[str, Any]], Optional[str]]] = None,
) -> Document:
    """
    Build a .docx file from a JSON object (or path) produced by docx2json
    — tolerant of JSON modified by an LLM.

    Args:
        json_source: dict or path to JSON file.
        output_docx_path: where to save the resulting .docx.
        resources_dir: optional directory to resolve embedded image filenames.
        image_resolver: optional callable(images_info_dict) -> absolute path;
                        used to locate original images for re-embedding.

    Returns:
        The python-docx Document object (already saved to output_docx_path).
    """
    if isinstance(json_source, str):
        with open(json_source, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json_source

    doc = Document()

    # Meta / core properties
    meta = data.get("meta") or {}
    core_props = meta.get("core_properties") or {}
    _apply_core_properties(doc, core_props)

    # Sections (apply first section only)
    _apply_section_settings(doc, data.get("sections") or [])

    # Blocks
    for block in data.get("blocks", []):
        btype = block.get("type")
        if btype == "paragraph":
            _write_paragraph(doc, block, image_resolver, resources_dir)
        elif btype == "table":
            _write_table(doc, block, image_resolver, resources_dir)
        else:
            # Unknown block type: ignore
            pass

    # Save
    os.makedirs(os.path.dirname(os.path.abspath(output_docx_path)), exist_ok=True)
    doc.save(output_docx_path)
    return doc

# Example usage:
# doc = json2docx("output.json", "reconstructed.docx", resources_dir="images/")
# Or with a dict:
# data = docx2json("input.docx")
# doc = json2docx(data, "roundtrip.docx")

# def main():
#     """Test function to generate both normal and compact JSON."""
#     docx_input = "test1.docx"
#     json_output_normal = "output_normal.json"
#     json_output_compact = "output_compact.json"
#
#     print(f"Generating normal JSON from {docx_input}...")
#     docx2json(docx_input, json_path=json_output_normal, compact=False)
#     print(f"Normal JSON saved to {json_output_normal}")
#
#     print(f"Generating compact JSON from {docx_input}...")
#     docx2json(docx_input, json_path=json_output_compact, compact=True)
#     print(f"Compact JSON saved to {json_output_compact}")
#
# if __name__ == "__main__":
#     main()
